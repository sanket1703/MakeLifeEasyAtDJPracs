from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

email_my = 'liberate.from.mu@gmail.com'
password = 'djsce@12345'
message = 'This is my message'
##########################################################
###   ALWAYS CHECK ALL DETAILS BEFORE SENDING   ##########
##########################################################
Subject = ' OSTL Labs'
Labs = 'OSTL'
#Main Experiment for Folder
Experiment = '6'
#Sub-Experiment inside Folder
Experiment_No = '6'
#Original Directory
Folder_Name = 'Experiments' 
#Subject in each file
subject = Labs + ' - Experiment ' + Experiment_No
#File from where we want to make copies
File_name = 'p6'
############################################################
############################################################
############################################################
msg = MIMEMultipart()
msg.attach(MIMEText(message, 'plain'))


name = ['Shrey D.','Samit K.','Saurav T.','Sheerang T.','Sakshi P.','Samiksha P.','Saakshi S.','Rutu G.','Sharayu B.','Sakshi U.','Shubh Nandu','Shruti S.','Shivesh J.','Sanay S.','Sanket S.','Siddharth T.','Sarvesh N.']
sap = ['60004180102','60004180093','60004180097','60004180101', '60004180091', '60004180092', '60004180089', '60004180087', '60004180098', '60004180090', '60004180104', '60004180103', '60004180100', '60004180094', '60004180095','60004180106','60004180096']
email = ['shreydedhia@yahoo.in','samitkk18@gmail.com','sauravtiwari27@gmail.com','shreerang2604@gmail.com', 'sakshi16.patil2000@gmail.com', 'samiksha16.patil@gmail.com', 'shahsaakshi25@gmail.com', 'rutugaglani@gmail.com', 'bokde.sharayu@gmail.com', 'sakshiuppoor@gmail.com', 'shubh.b.nandu@gmail.com', 'sawantshruti334@gmail.com', 'shubjha12@gmail.com', 'sanayshah2@gmail.com', 'sanketyou8@gmail.com','siddharthtrivedi19@gmail.com','sarveshnavare7@gmail.com']
#email  = ['sanketyou8@gmail.com']
#sap = ['60004180095']
#name = ['Sanket S']

dec = input('Do you want to send e-mails?')
if dec == 'y':
    flag = 1
else:
    flag = 0 

if os.path.exists(Folder_Name):
    print('path exists')
    os.chdir(Folder_Name)
else:
    os.mkdir(Folder_Name)
    print('made folder')
    os.chdir(Folder_Name)
    
server = smtplib.SMTP('smtp.gmail.com', 587)
server.starttls()
server.login(email_my, password)

for i in range(len(name)):
    #Change Experiment Number
 
    info = 'Name : '+name[i]+'\nSAP : '+sap[i]+'\nSubject: '+subject+'\n'
    j= 'Exp_'+Experiment_No+'_'+sap[i]+'_'+name[i]
    if os.path.exists(sap[i]):
        print('SAP path exists')
        os.chdir(sap[i])
    else:
        os.mkdir(sap[i])
        os.chdir(sap[i])
    
    if os.path.exists('Experiment_'+Experiment):
        print('Exp path exists')
        os.chdir('Experiment_'+Experiment)
    else:
        os.mkdir('Experiment_'+Experiment)
        os.chdir('Experiment_'+Experiment)
    # For normal File operation
    
    # f = open(j+'.txt','w')
    # dst = j+'.txt'
    # #Change File Name
    # f1 = open('/Users/apple/Desktop/Division_B2/'+File_name+'.txt','r') 
    
    # r = f1.read()
    # # shutil.copy('/Users/apple/Desktop/untitled folder/test.txt',dst)
    # f.write(info+'\n\n\n'+r)
    # f1.close()
    # f.close()

    base = Document()
    dst = j+'.docx'
    #stuff to be added in the base
    para=base.add_paragraph(info+'\n\n\n')
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    composer = Composer(base)
    added = Document('/Users/apple/Desktop/Division_B2/'+File_name +'.docx')

    composer.append(added)
    composer.save(j+'.docx')

    print('Saved')

    
    #Change if using a Dictionary
    send_to_email = email[i]
    #Change File Location
    #file_location = '/Users/apple/Desktop/Division_B2/'+Folder_Name+'/'+sap[i]+'/'+j+'.txt'
    file_location = '/Users/apple/Desktop/Division_B2/'+Folder_Name+'/'+sap[i]+'/'+'Experiment_'+Experiment+'/'+j+'.docx'

    msg = MIMEMultipart()
    msg['From'] = email_my
    msg['To'] = email[i]
    #Change subject here
    msg['Subject'] = Subject
    #body = '''POSTLABS for Experiment 1 '''
    
    filename = os.path.basename(file_location)
    attachment = open(file_location, "rb")
    part = MIMEBase('application', 'octet-stream')
    part.set_payload(attachment.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', "attachment; filename= %s" % filename)
    msg.attach(part)
    if flag is 1:
        server.login(email_my, password)
        text = msg.as_string()

        server.sendmail(email_my, email[i], text)
        print("sent")
    os.chdir('..')
    os.chdir('..')
    
server.close()
print('Server Close')


